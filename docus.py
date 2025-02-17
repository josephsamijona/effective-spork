from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Changé de WP_ALIGN_PARAGRAPH à WD_ALIGN_PARAGRAPH
import svglib.svglib
from reportlab.graphics import renderPM
import io

def create_wave_svg(width=800, height=200, is_header=True):
    """
    Crée un design de vague en SVG
    """
    svg_template = f'''<?xml version="1.0" encoding="UTF-8"?>
    <svg width="{width}" height="{height}" viewBox="0 0 {width} {height}" xmlns="http://www.w3.org/2000/svg">
        <defs>
            <linearGradient id="wave-gradient" x1="0%" y1="0%" x2="100%" y2="0%">
                <stop offset="0%" style="stop-color:#0047AB;stop-opacity:1" />
                <stop offset="100%" style="stop-color:#4169E1;stop-opacity:1" />
            </linearGradient>
        </defs>
        <path d="M0 {height if not is_header else 0}
                 C {width/4} {height * (0.7 if is_header else 0.3)},
                   {width/2} {height * (0.3 if is_header else 0.7)},
                   {width} {height if is_header else 0}
                 L{width} {height if not is_header else 0} L0 {height if not is_header else 0} Z"
              fill="url(#wave-gradient)"/>
    </svg>'''
    return svg_template

def svg_to_png(svg_content):
    """
    Convertit le SVG en PNG pour l'insertion dans le document
    """
    drawing = svglib.svglib.svg2rlg(io.StringIO(svg_content))
    return renderPM.drawToString(drawing, fmt='PNG')

def create_document_template():
    # Créer un nouveau document
    doc = Document()
    
    # Configurer les marges du document
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # Créer et ajouter le header avec la vague
    header_svg = create_wave_svg(is_header=True)
    header_png = svg_to_png(header_svg)
    
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Utilise WD_ALIGN_PARAGRAPH au lieu de WP_ALIGN_PARAGRAPH
    run = header_para.add_run()
    run.add_picture(io.BytesIO(header_png), width=Inches(8.27))
    
    # Ajouter du contenu exemple
    doc.add_heading('Titre du Document', 0)
    doc.add_paragraph('Contenu du document...')
    
    # Créer et ajouter le footer avec la vague
    footer_svg = create_wave_svg(is_header=False)
    footer_png = svg_to_png(footer_svg)
    
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    run.add_picture(io.BytesIO(footer_png), width=Inches(8.27))
    
    # Sauvegarder le document
    doc.save('template_with_waves.docx')

if __name__ == "__main__":
    create_document_template()